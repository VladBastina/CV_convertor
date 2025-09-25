import os, json
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image

# ----------------- CONFIG -----------------
SVG_PATH = "zega_logo.svg"
CONVERTED_PNG = "profile_converted.png"
OUTPUT_DOCX = "Denisa_CV_from_json.docx"
IMAGE_WIDTH_INCHES = 1.5
# ------------------------------------------


def set_run_font(run, name="Calibri", size_pt=10.5, bold=False, italic=False, color=(0, 0, 0)):
    run.font.name = name
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor(*color)
    try:
        r = run._element.rPr.rFonts
        r.set(qn("w:eastAsia"), name)
    except Exception:
        pass


def add_section_heading(doc, text, color=(0, 112, 192)):
    p = doc.add_paragraph()
    run = p.add_run(text.upper())
    set_run_font(run, size_pt=14, bold=True, color=color)
    p.space_after = Pt(6)
    return p


def add_shaded_spacer(doc, height_pt=6, color="D9D9D9"):
    p = doc.add_paragraph()
    p.add_run("")
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color)
    pPr.append(shd)
    p.paragraph_format.space_before = Pt(height_pt)
    p.paragraph_format.space_after = Pt(height_pt)


# ---------- Section Renderers ----------
def render_bullets(doc, section):
    for item in section.get("items", []):
        pb = doc.add_paragraph(style="List Bullet")
        if "label" in item:
            run = pb.add_run(f"{item['label']}: ")
            set_run_font(run, bold=True)
            run2 = pb.add_run(item.get("value", ""))
            set_run_font(run2)
        else:
            run = pb.add_run(item.get("value", ""))
            set_run_font(run)


def render_experience(doc, section):
    for pos in section.get("positions", []):
        p = doc.add_paragraph()
        run = p.add_run(pos["role"])
        set_run_font(run, bold=True)
        if "dates" in pos:
            run2 = p.add_run(f"    {pos['dates']}")
            set_run_font(run2, size_pt=10, italic=True, color=(89, 89, 89))
        for d in pos.get("details", []):
            bl = doc.add_paragraph(style="List Bullet")
            r = bl.add_run(d)
            set_run_font(r)
        if "technologies" in pos:
            p = doc.add_paragraph()
            run = p.add_run("Technologies: ")
            set_run_font(run, bold=True, size_pt=10)
            run2 = p.add_run(pos["technologies"])
            set_run_font(run2, size_pt=10, color=(89, 89, 89))


def render_list(doc, section):
    for c in section.get("items", []):
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(c)
        set_run_font(run)


def render_text(doc, section):
    for item in section.get("items", []):
        p = doc.add_paragraph()
        run = p.add_run(item["value"])
        set_run_font(run, bold=True)
        if "sub" in item:
            run2 = p.add_run("    " + item["sub"])
            set_run_font(run2, size_pt=10, italic=True, color=(89, 89, 89))


# Mapping section type → renderer
SECTION_RENDERERS = {
    "bullets": render_bullets,
    "experience": render_experience,
    "list": render_list,
    "text": render_text,
}


# ---------- Main builder ----------
def create_cv_docx(data: dict):
    # Convert PNG if available
    if os.path.exists(CONVERTED_PNG):
        im = Image.open(CONVERTED_PNG)
        im.convert("RGBA").save(CONVERTED_PNG)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    style.font.size = Pt(11)

    # --- HEADER ---
    header = doc.add_table(rows=1, cols=2)
    header.columns[0].width = Inches(5.5)
    header.columns[1].width = Inches(1.5)

    left_cell = header.cell(0, 0)
    right_cell = header.cell(0, 1)

    # Name
    p = left_cell.paragraphs[0]
    run = p.add_run(data.get("name", ""))
    set_run_font(run, size_pt=24, bold=True)

    # Role
    if "role" in data:
        p = left_cell.add_paragraph()
        run = p.add_run(data["role"])
        set_run_font(run, size_pt=12, color=(89, 89, 89))

    # Photo
    if os.path.exists(CONVERTED_PNG):
        try:
            p = right_cell.paragraphs[0]
            p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            r = p.add_run()
            r.add_picture(CONVERTED_PNG, width=Inches(IMAGE_WIDTH_INCHES))
        except Exception:
            pass

    # --- ABOUT ---
    if "about" in data:
        p = doc.add_paragraph(data["about"])
        set_run_font(p.runs[0])
        add_shaded_spacer(doc)

    # --- SECTIONS ---
    for section in data.get("sections", []):
        add_section_heading(doc, section.get("title", ""))
        renderer = SECTION_RENDERERS.get(section.get("type"))
        if renderer:
            renderer(doc, section)
        else:
            # fallback: dump text if unknown section type
            for line in section.get("items", []):
                p = doc.add_paragraph(str(line))
                set_run_font(p.add_run(str(line)))

    # Save
    doc.save(OUTPUT_DOCX)
    print(f"Saved CV to: {OUTPUT_DOCX}")
